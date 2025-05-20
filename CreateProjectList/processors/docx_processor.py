from pathlib import Path
from win32com import client
import pythoncom
from docx import Document
from typing import Dict
from CreateProjectList.processors.document_processor_base import DocumentProcessorBase
from CreateProjectList.utils.log_manager import LogManager

class DOCXProcessor(DocumentProcessorBase):
    """DOCX形式のWordファイル処理に特化したプロセッサー"""
    def __init__(self):
        """初期化"""
        self.logger = LogManager().get_logger(self.__class__.__name__)
    
    def can_process(self, file_path: Path) -> bool:
        """
        指定されたファイルが処理可能か判定
        
        Args:
            file_path: 処理対象ファイルのパス
            
        Returns:
            bool: docxファイルの場合True
        """
        return file_path.suffix.lower() == '.docx'

    def process_file(self, input_path: Path, output_path: Path, replacements: Dict[str, str]) -> None:
        """
        DOCXファイルを処理
        
        Args:
            input_path: 入力ファイルのパス
            output_path: 出力ファイルのパス
            replacements: 置換ルール辞書
        """
        word = None
        document = None

        try:
            pythoncom.CoInitialize()
            
            self._report_progress(0, "DOCXファイルを処理中...", str(input_path.name))

            # Wordアプリケーションの起動（DOCX用の最適化設定）
            word = client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            # DOCX互換モードの設定
            word.Options.ConfirmConversions = False
            
            self._report_progress(20, "ファイルを開いています...", str(input_path.name))
            
            # 修復モードでファイルを開く
            try:
                document = Document(input_path.resolve())
            except:
                document = Document(input_path.resolve())
            
            # 本文処理
            self._report_progress(40, "本文を処理中...", "")
            self._process_content(document, replacements)
            
            # 図形処理
            self._report_progress(60, "図形を処理中...", "")
            self._process_shapes(document, replacements)
            
            # DOCX形式で保存（最新フォーマット）
            document.save(str(output_path.resolve()))
            
            self._report_progress(100, "処理完了", f"{input_path.name} -> {output_path.name}")

        except Exception as e:
            self.logger.error(f"DOCX処理エラー - {input_path}: {str(e)}")
            raise
        
        finally:
            if word:
                word.Quit()
            pythoncom.CoUninitialize()

    def _process_content(self, document, replacements: Dict[str, str]) -> None:
        """
        本文テキストおよびテーブルセルのテキストを処理（DOC用）

        Args:
            document: 処理対象ドキュメント
            replacements: 置換ルール辞書
        """
        try:
            # 段落内のテキストを置換
            for paragraph in document.paragraphs:
                for old_text, new_text in replacements.items():
                    if old_text in paragraph.text:
                        # 「なし」「未設定」の場合は空文字に置換
                        replacement_text = "" if new_text in ["なし", "未設定"] else new_text
                        for run in paragraph.runs:
                            run.text = run.text.replace(old_text, replacement_text)

            # テーブル内のテキストを置換
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for old_text, new_text in replacements.items():
                                if old_text in paragraph.text:
                                    # 「なし」「未設定」の場合は空文字に置換
                                    replacement_text = "" if new_text in ["なし", "未設定"] else new_text
                                    for run in paragraph.runs:
                                        run.text = run.text.replace(old_text, replacement_text)

        except Exception as e:
            print(f"置換処理中にエラーが発生しました: {str(e)}")


    def _process_shapes(self, document, replacements: Dict[str, str]) -> None:
        """
        Word文書内の全図形のテキストを置換する関数
        
        Args:
            doc (Document): 処理対象のDocument オブジェクト
            replacements (Dict[str, str]): 置換対象の文字列辞書
        
        Returns:
            Document: 処理済みのDocument オブジェクト
        """
        try:
            # インライン図形の処理
            for shape in document.inline_shapes:
                if hasattr(shape, '_inline'):
                    shape_element = shape._inline.graphic.graphicData
                    txbx_elements = shape_element.findall('.//w:txbxContent', 
                                                        {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    
                    for txbx in txbx_elements:
                        for paragraph in txbx.findall('.//w:p', 
                                                    {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            for run in paragraph.findall('.//w:t', 
                                                    {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                if run.text:
                                    text = run.text
                                    for old_text, new_text in replacements.items():
                                        # 「なし」「未設定」の場合は空文字に置換
                                        replacement_text = "" if new_text in ["なし", "未設定"] else new_text
                                        text = text.replace(old_text, replacement_text)
                                    run.text = text
                    
            # 浮動図形の処理
            for shape in document._element.body.findall('.//w:drawing',
                                                {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                txbx_elements = shape.findall('.//w:txbxContent', 
                                            {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                
                for txbx in txbx_elements:
                    for paragraph in txbx.findall('.//w:p', 
                                                {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        for run in paragraph.findall('.//w:t', 
                                                {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            if run.text:
                                text = run.text
                                for old_text, new_text in replacements.items():
                                    # 「なし」「未設定」の場合は空文字に置換
                                    replacement_text = "" if new_text in ["なし", "未設定"] else new_text
                                    text = text.replace(old_text, replacement_text)
                                run.text = text

        except Exception as e:
            shape_name = getattr(shape, 'Name', 'unknown')
            self.logger.warning(f"DOCX Shape '{shape_name}' の処理でエラー: {str(e)}")