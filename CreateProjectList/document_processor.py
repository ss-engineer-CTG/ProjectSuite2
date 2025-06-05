"""
統合ドキュメント処理エンジン
全ファイル形式の処理を統合管理
"""
import logging
from pathlib import Path
from typing import Dict, List, Optional, Callable, Any, Tuple
import tempfile
import shutil
import re
from datetime import datetime

# Office関連の条件付きインポート
try:
    from win32com import client
    import pythoncom
    WIN32_AVAILABLE = True
except ImportError:
    WIN32_AVAILABLE = False
    logging.warning("win32com が利用できません。Office系ファイル処理に制限があります")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

class DocumentProcessor:
    """統合ドキュメント処理クラス"""
    
    # 対応ファイル形式
    SUPPORTED_EXTENSIONS = {
        '.doc': 'word_old',
        '.docx': 'word_new', 
        '.xls': 'excel_old',
        '.xlsx': 'excel_new',
        '.xlsm': 'excel_macro'
    }
    
    def __init__(self, core_manager):
        """
        初期化
        
        Args:
            core_manager: CoreManagerインスタンス
        """
        self.core_manager = core_manager
        self.logger = logging.getLogger(self.__class__.__name__)
        
        # 処理状態
        self.current_project_data: Optional[Dict[str, Any]] = None
        self.is_db_connected = False
        self.temp_dir: Optional[Path] = None
        
        # コールバック
        self._progress_callback: Optional[Callable] = None
        self._cancel_check: Optional[Callable] = None
        
        # 初期化
        self._initialize_temp_dir()
        
    def _initialize_temp_dir(self):
        """一時ディレクトリの初期化"""
        try:
            temp_base = self.core_manager.get_temp_dir()
            if temp_base:
                self.temp_dir = Path(temp_base)
                self.temp_dir.mkdir(parents=True, exist_ok=True)
            else:
                self.temp_dir = Path(tempfile.mkdtemp(prefix="doc_processor_"))
            
            self.logger.info(f"一時ディレクトリを初期化: {self.temp_dir}")
            
        except Exception as e:
            self.logger.error(f"一時ディレクトリ初期化エラー: {e}")
            self.temp_dir = Path(tempfile.mkdtemp(prefix="doc_processor_"))
    
    def connect_database(self) -> bool:
        """データベース接続"""
        try:
            self.is_db_connected = self.core_manager.test_database_connection()
            if self.is_db_connected:
                self.logger.info("データベース接続成功")
            else:
                self.logger.error("データベース接続失敗")
            return self.is_db_connected
        except Exception as e:
            self.logger.error(f"データベース接続エラー: {e}")
            return False
    
    def fetch_project_data(self, project_id: int) -> Optional[Dict[str, Any]]:
        """プロジェクトデータの取得"""
        try:
            return self.core_manager.get_project_data(project_id)
        except Exception as e:
            self.logger.error(f"プロジェクトデータ取得エラー: {e}")
            return None
    
    def get_all_projects(self) -> List[Dict[str, Any]]:
        """全プロジェクトの取得"""
        try:
            return self.core_manager.get_all_projects()
        except Exception as e:
            self.logger.error(f"プロジェクト一覧取得エラー: {e}")
            return []
    
    def set_project_data(self, project_data: Dict[str, Any]):
        """プロジェクトデータの設定"""
        self.current_project_data = project_data
        self.logger.info(f"プロジェクトデータを設定: {project_data.get('project_name', 'Unknown')}")
    
    def can_process_file(self, file_path: Path) -> bool:
        """ファイルが処理可能か判定"""
        return file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS
    
    def process_documents(self, input_folder_path: str, output_folder_path: str,
                         progress_callback: Optional[Callable] = None,
                         cancel_check: Optional[Callable] = None) -> Dict:
        """
        フォルダ内のドキュメントを一括処理
        
        Args:
            input_folder_path: 入力フォルダパス
            output_folder_path: 出力フォルダパス
            progress_callback: 進捗コールバック
            cancel_check: キャンセルチェックコールバック
            
        Returns:
            Dict: 処理結果
        """
        if not self.current_project_data:
            raise ValueError("プロジェクトが選択されていません")
        
        # コールバック設定
        self._progress_callback = progress_callback
        self._cancel_check = cancel_check
        
        input_path = Path(input_folder_path)
        output_path = Path(output_folder_path)
        
        try:
            self._report_progress(0, "処理を開始しています...")
            
            # 入力検証
            self._validate_input(input_path, output_path)
            
            # 置換ルール作成
            replacements = self._create_replacements()
            
            # 処理対象ファイル取得
            target_files = self._get_target_files(input_path)
            if not target_files:
                return {'processed': [], 'errors': [], 'cancelled': False}
            
            # フォルダ構造作成
            self._report_progress(10, "フォルダ構造を作成中...")
            self._create_folder_structure(input_path, output_path, replacements)
            
            # ファイル処理
            return self._process_files(target_files, input_path, output_path, replacements)
            
        except Exception as e:
            self.logger.error(f"ドキュメント処理エラー: {e}")
            raise
        finally:
            self._cleanup_temp_files()
    
    def _validate_input(self, input_path: Path, output_path: Path):
        """入力検証"""
        if not input_path.exists():
            raise ValueError("入力フォルダが存在しません")
        if not input_path.is_dir():
            raise ValueError("入力パスがフォルダではありません")
        if not any(input_path.iterdir()):
            raise ValueError("入力フォルダが空です")
        
        # 出力フォルダ作成
        try:
            output_path.mkdir(parents=True, exist_ok=True)
        except Exception as e:
            raise ValueError(f"出力フォルダの作成に失敗: {e}")
    
    def _create_replacements(self) -> Dict[str, str]:
        """置換ルール辞書の作成"""
        replacements = {}
        rules = self.core_manager.get_replacement_rules()
        
        for rule in rules:
            key = rule.get('search', '')
            value_key = rule.get('replace', '')
            value = str(self.current_project_data.get(value_key, ''))
            
            # 空値の処理
            if value.lower() in ['なし', 'none', '']:
                value = ''
            
            replacements[key] = value
        
        return replacements
    
    def _get_target_files(self, folder_path: Path) -> List[Path]:
        """処理対象ファイルの取得"""
        files = []
        for file_path in folder_path.rglob('*'):
            if file_path.is_file() and self.can_process_file(file_path):
                files.append(file_path)
        return files
    
    def _create_folder_structure(self, input_path: Path, output_path: Path, replacements: Dict[str, str]):
        """フォルダ構造の作成"""
        folders = [f for f in input_path.rglob('*') if f.is_dir()]
        
        for folder in folders:
            relative_path = folder.relative_to(input_path)
            processed_path = self._process_path(relative_path, replacements)
            target_path = output_path / processed_path
            target_path.mkdir(parents=True, exist_ok=True)
    
    def _process_path(self, path: Path, replacements: Dict[str, str]) -> Path:
        """パスの処理（フォルダ名置換）"""
        if str(path) == '.':
            return path
        
        parts = []
        for part in path.parts:
            processed_part = self._process_text(part, replacements)
            processed_part = self._sanitize_filename(processed_part)
            parts.append(processed_part)
        
        return Path(*parts)
    
    def _process_files(self, files: List[Path], input_root: Path, output_root: Path, 
                      replacements: Dict[str, str]) -> Dict:
        """ファイル一括処理"""
        processed_files = []
        errors = []
        total_files = len(files)
        
        for i, file_path in enumerate(files, 1):
            if self._should_cancel():
                return {'processed': processed_files, 'errors': errors, 'cancelled': True}
            
            try:
                progress = 20 + ((i / total_files) * 70)
                self._report_progress(progress, f"ファイル処理中 ({i}/{total_files})", file_path.name)
                
                self._process_single_file(file_path, input_root, output_root, replacements)
                processed_files.append(file_path)
                
            except Exception as e:
                self.logger.error(f"ファイル処理エラー {file_path}: {e}")
                errors.append((file_path, str(e)))
        
        self._report_progress(100, "処理完了")
        return {'processed': processed_files, 'errors': errors, 'cancelled': False}
    
    def _process_single_file(self, file_path: Path, input_root: Path, output_root: Path,
                           replacements: Dict[str, str]):
        """単一ファイル処理"""
        # 出力パス生成
        relative_path = file_path.relative_to(input_root)
        processed_path = self._process_path(relative_path.parent, replacements)
        processed_name = self._process_text(relative_path.name, replacements)
        processed_name = self._sanitize_filename(processed_name)
        
        output_path = output_root / processed_path / processed_name
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # ファイル形式に応じた処理
        file_type = self.SUPPORTED_EXTENSIONS[file_path.suffix.lower()]
        
        if file_type in ['word_old', 'word_new']:
            self._process_word_file(file_path, output_path, replacements)
        elif file_type in ['excel_old', 'excel_new', 'excel_macro']:
            self._process_excel_file(file_path, output_path, replacements)
        else:
            raise ValueError(f"未対応のファイル形式: {file_path.suffix}")
    
    def _process_word_file(self, input_path: Path, output_path: Path, replacements: Dict[str, str]):
        """Word ファイル処理"""
        if input_path.suffix.lower() == '.docx' and DOCX_AVAILABLE:
            self._process_docx_with_docx_lib(input_path, output_path, replacements)
        elif WIN32_AVAILABLE:
            self._process_word_with_com(input_path, output_path, replacements)
        else:
            raise RuntimeError("Word ファイル処理に必要なライブラリが不足しています")
    
    def _process_excel_file(self, input_path: Path, output_path: Path, replacements: Dict[str, str]):
        """Excel ファイル処理"""
        if input_path.suffix.lower() in ['.xlsx', '.xlsm'] and OPENPYXL_AVAILABLE:
            self._process_xlsx_with_openpyxl(input_path, output_path, replacements)
        elif WIN32_AVAILABLE:
            self._process_excel_with_com(input_path, output_path, replacements)
        else:
            raise RuntimeError("Excel ファイル処理に必要なライブラリが不足しています")
    
    def _process_docx_with_docx_lib(self, input_path: Path, output_path: Path, replacements: Dict[str, str]):
        """python-docx を使用した DOCX 処理"""
        try:
            doc = DocxDocument(str(input_path))
            
            # 段落のテキスト置換
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run.text = self._process_text(run.text, replacements)
            
            # テーブルのテキスト置換
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.text = self._process_text(run.text, replacements)
            
            doc.save(str(output_path))
            
        except Exception as e:
            raise RuntimeError(f"DOCX処理エラー: {e}")
    
    def _process_xlsx_with_openpyxl(self, input_path: Path, output_path: Path, replacements: Dict[str, str]):
        """openpyxl を使用した XLSX 処理"""
        try:
            wb = openpyxl.load_workbook(str(input_path))
            
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            cell.value = self._process_text(cell.value, replacements)
            
            wb.save(str(output_path))
            
        except Exception as e:
            raise RuntimeError(f"XLSX処理エラー: {e}")
    
    def _process_word_with_com(self, input_path: Path, output_path: Path, replacements: Dict[str, str]):
        """COM を使用した Word 処理"""
        word = None
        document = None
        
        try:
            pythoncom.CoInitialize()
            
            word = client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False
            
            document = word.Documents.Open(str(input_path.resolve()))
            
            # テキスト置換
            for old_text, new_text in replacements.items():
                find = document.Content.Find
                find.ClearFormatting()
                find.Replacement.ClearFormatting()
                find.Execute(
                    FindText=old_text,
                    ReplaceWith=new_text,
                    Replace=2,  # 全て置換
                    Forward=True,
                    Wrap=1
                )
            
            # 保存
            if input_path.suffix.lower() == '.doc':
                document.SaveAs(str(output_path.resolve()), FileFormat=0)  # .doc
            else:
                document.SaveAs(str(output_path.resolve()), FileFormat=16)  # .docx
            
        except Exception as e:
            raise RuntimeError(f"Word COM処理エラー: {e}")
        finally:
            if document:
                document.Close(SaveChanges=False)
            if word:
                word.Quit()
            pythoncom.CoUninitialize()
    
    def _process_excel_with_com(self, input_path: Path, output_path: Path, replacements: Dict[str, str]):
        """COM を使用した Excel 処理"""
        excel = None
        workbook = None
        
        try:
            pythoncom.CoInitialize()
            
            excel = client.Dispatch("Excel.Application")
            excel.Visible = False
            excel.DisplayAlerts = False
            
            workbook = excel.Workbooks.Open(str(input_path.resolve()))
            
            # 全シートでテキスト置換
            for sheet in workbook.Worksheets:
                for old_text, new_text in replacements.items():
                    found_cell = sheet.Cells.Find(
                        What=old_text,
                        LookAt=2,  # 部分一致
                        MatchCase=True
                    )
                    
                    if found_cell:
                        first_address = found_cell.Address
                        while True:
                            original_text = str(found_cell.Value)
                            new_value = self._process_text(original_text, {old_text: new_text})
                            if original_text != new_value:
                                found_cell.Value = new_value
                            
                            found_cell = sheet.Cells.FindNext(found_cell)
                            if not found_cell or found_cell.Address == first_address:
                                break
            
            # 保存
            if input_path.suffix.lower() == '.xls':
                workbook.SaveAs(str(output_path.resolve()), FileFormat=56)  # .xls
            elif input_path.suffix.lower() == '.xlsm':
                workbook.SaveAs(str(output_path.resolve()), FileFormat=52)  # .xlsm
            else:
                workbook.SaveAs(str(output_path.resolve()), FileFormat=51)  # .xlsx
            
        except Exception as e:
            raise RuntimeError(f"Excel COM処理エラー: {e}")
        finally:
            if workbook:
                workbook.Close(SaveChanges=False)
            if excel:
                excel.Quit()
            pythoncom.CoUninitialize()
    
    def _process_text(self, text: str, replacements: Dict[str, str]) -> str:
        """テキスト置換処理"""
        if not isinstance(text, str):
            return text
        
        result = text
        for old_text, new_text in replacements.items():
            if str(new_text).lower() in ['なし', 'none']:
                result = result.replace(old_text, '')
            else:
                result = result.replace(old_text, str(new_text))
        
        return result
    
    def _sanitize_filename(self, name: str) -> str:
        """ファイル名のサニタイズ"""
        # 不正文字を除去
        sanitized = re.sub(r'[<>:"/\\|?*]', '_', name)
        sanitized = sanitized.strip(' .')
        
        if not sanitized:
            sanitized = "unnamed"
        
        # 長さ制限
        if len(sanitized) > 255:
            sanitized = sanitized[:255]
        
        return sanitized
    
    def _report_progress(self, progress: float, status: str, detail: str = ""):
        """進捗報告"""
        if self._progress_callback:
            self._progress_callback(progress, status, detail)
    
    def _should_cancel(self) -> bool:
        """キャンセル確認"""
        return bool(self._cancel_check and self._cancel_check())
    
    def _cleanup_temp_files(self):
        """一時ファイルのクリーンアップ"""
        try:
            if self.temp_dir and self.temp_dir.exists():
                for item in self.temp_dir.glob('*'):
                    if item.is_file():
                        item.unlink()
                    elif item.is_dir():
                        shutil.rmtree(item)
        except Exception as e:
            self.logger.error(f"一時ファイルクリーンアップエラー: {e}")
    
    def __del__(self):
        """デストラクタ"""
        try:
            self._cleanup_temp_files()
        except:
            pass